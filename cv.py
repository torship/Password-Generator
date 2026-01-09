from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word document
doc = Document()

# Define styles
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False):
    run = doc.add_paragraph().add_run(text)
    run.bold = bold
    run.font.size = Pt(11)

# Header
doc.add_heading('SHOUVIK CHAND', level=0)
doc.add_paragraph('📍 Durgapur, West Bengal, India')
doc.add_paragraph('📞 +91-9883151363 | ✉️ shouvikchand0@gmail.com')

# Professional Summary
add_heading('PROFESSIONAL SUMMARY', level=1)
doc.add_paragraph(
    "Dedicated and self-motivated BCA graduate from Dr. B.C. Roy Academy with hands-on experience in "
    "Node.js, Express.js, and front-end technologies like HTML, CSS, and JavaScript. Passionate about "
    "building practical web applications and eager to contribute to dynamic development teams. Strong "
    "foundation in computer science with a commitment to continuous learning."
)

# Technical Skills
add_heading('TECHNICAL SKILLS', level=1)
doc.add_paragraph("Backend: Node.js, Express.js")
doc.add_paragraph("Frontend: HTML, CSS, JavaScript")
doc.add_paragraph("Programming: C Language")
doc.add_paragraph("Tools/Technologies: Git, VS Code, REST APIs")

# Education
add_heading('EDUCATION', level=1)
doc.add_paragraph("Bachelor of Computer Applications (BCA)", style='List Bullet')
doc.add_paragraph("Dr. B.C. Roy Academy of Professional Courses, Durgapur")
doc.add_paragraph("2020 – 2023 | 68.6%")
doc.add_paragraph("Class XII – West Bengal Board", style='List Bullet')
doc.add_paragraph("Medium: Bengali | Year: 2023 | 68.6%")
doc.add_paragraph("Class X – West Bengal Board", style='List Bullet')
doc.add_paragraph("Medium: Bengali | Year: 2021 | 69%")

# Projects
add_heading('PROJECTS', level=1)
doc.add_paragraph("Weather App", style='List Bullet')
doc.add_paragraph("April 2025 – April 2025")
doc.add_paragraph("A Node.js application that provides real-time weather reports based on city input using APIs. "
                  "Demonstrates backend logic integration with real-time data fetching.")

# Languages
add_heading('LANGUAGES', level=1)
doc.add_paragraph("Bengali: Read, Write, Speak")
doc.add_paragraph("English: Read, Write, Speak")
doc.add_paragraph("Hindi: Speak")

# Personal Details
add_heading('PERSONAL DETAILS', level=1)
doc.add_paragraph("Date of Birth: September 7, 2005")
doc.add_paragraph("Gender: Male")

# Save the document
output_path = "/mnt/data/Shouvik_Chand_CV.docx"
doc.save(output_path)

output_path
